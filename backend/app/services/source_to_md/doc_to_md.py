"""
Word/DOCX 转 Markdown 解析器
提取 DOCX 文本、表格、图片，保留层级结构
"""

from pathlib import Path
from typing import List


def parse_docx(file_path: str) -> str:
    """
    解析 DOCX 文件，返回 Markdown 格式文本
    
    Args:
        file_path: DOCX 文件路径
    
    Returns:
        Markdown 格式的文本内容
    """
    try:
        from docx import Document
        from docx.shared import Pt
        
        doc = Document(file_path)
        markdown_lines = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # 根据样式判断标题层级
            style_name = para.style.name.lower() if para.style else ""
            
            if 'heading 1' in style_name or '标题 1' in style_name:
                markdown_lines.append(f"# {text}")
            elif 'heading 2' in style_name or '标题 2' in style_name:
                markdown_lines.append(f"## {text}")
            elif 'heading 3' in style_name or '标题 3' in style_name:
                markdown_lines.append(f"### {text}")
            elif 'heading 4' in style_name or '标题 4' in style_name:
                markdown_lines.append(f"#### {text}")
            else:
                # 正文
                # 处理加粗、斜体
                text_parts = []
                for run in para.runs:
                    run_text = run.text
                    if run.bold and run.italic:
                        run_text = f"***{run_text}***"
                    elif run.bold:
                        run_text = f"**{run_text}**"
                    elif run.italic:
                        run_text = f"*{run_text}*"
                    text_parts.append(run_text)
                
                para_text = ''.join(text_parts)
                markdown_lines.append(para_text)
        
        # 处理表格
        for table in doc.tables:
            markdown_lines.append("\n")
            # 表头
            if table.rows:
                header_cells = [cell.text.strip() for cell in table.rows[0].cells]
                markdown_lines.append("| " + " | ".join(header_cells) + " |")
                markdown_lines.append("| " + " | ".join(["---"] * len(header_cells)) + " |")
                
                # 数据行
                for row in table.rows[1:]:
                    cells = [cell.text.strip() for cell in row.cells]
                    markdown_lines.append("| " + " | ".join(cells) + " |")
            
            markdown_lines.append("\n")
        
        return '\n'.join(markdown_lines)
    
    except ImportError:
        return "# DOCX 解析失败\n\n请安装 python-docx：\n```bash\npip install python-docx\n```"
    
    except Exception as e:
        return f"# DOCX 解析错误\n\n错误信息：{str(e)}"


def extract_docx_images(file_path: str, output_dir: str) -> List[str]:
    """
    提取 DOCX 中的图片
    
    Args:
        file_path: DOCX 文件路径
        output_dir: 图片输出目录
    
    Returns:
        提取的图片路径列表
    """
    try:
        from docx import Document
        import zipfile
        import os
        
        doc = Document(file_path)
        output_path = Path(output_dir)
        output_path.mkdir(exist_ok=True)
        
        image_paths = []
        
        # DOCX 本质是 ZIP，从中提取图片
        with zipfile.ZipFile(file_path) as zip_file:
            for name in zip_file.namelist():
                if name.startswith('word/media/'):
                    img_data = zip_file.read(name)
                    img_name = Path(name).name
                    img_path = output_path / img_name
                    
                    with open(img_path, 'wb') as f:
                        f.write(img_data)
                    
                    image_paths.append(str(img_path))
        
        return image_paths
    
    except Exception as e:
        print(f"提取 DOCX 图片失败: {e}")
        return []
